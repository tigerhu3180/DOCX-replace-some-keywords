from docx import Document
from win32com import client
import os

OLDPATH = "C:\\Users\\qq\\Desktop\\doc\\old"
PATH = "C:\\Users\\qq\\Desktop\\doc\\new"

DICT = {
  "郭锋": "郭鑫",
  "安昊": "丁泽",
  "刘晓霖": "邢伟岭",
  "李扬": "鲍鹏举",
  "付晓蓉": "马竞楠",
  "刘宇梅": "郭烨红",
  "韩士准": "李莹",
  "刘达志": "赵宇",
  "赵飞": "王宇宁",
  "原伟": "陈晓宇",
  "刘改琴": "冯鑫",
  "2016":"2018",
  "2015":"2017",
  "TyCloud云管理平台开发项目（三期）":"MIMS通信机房维护巡检管理系统项目（一期）",
  "河北天翼科贸发展":"内蒙古电力集团蒙电信息通信产业",
  "有限公司":"有限责任公司",
  "TyCloud":"MIMS"
}

def main():
    for fileName in os.listdir(OLDPATH):
        if '.doc' in fileName and '.docx' not in fileName:
            # turn .doc to .docx and save it into the 'new' foler
            word = client.Dispatch('Word.Application')
            print(OLDPATH + "\\" + fileName)
            doc = word.Documents.Open(OLDPATH + "\\" + fileName)
            fileName = fileName.replace('.doc', '.docx') 
            doc.SaveAs(PATH + "\\" + fileName, 12)  
            doc.Close()
            word.Quit()

            oldFile = newFile = PATH + "\\" + fileName
        else:
            oldFile = OLDPATH + "\\" + fileName
            newFile = PATH + "\\" + fileName

        document = Document(oldFile)
        document = check(document)
        document.save(newFile)   

def check(document):
    # tables
    for table in document.tables:
        for row in range(len(table.rows)):
            for col in range(len(table.columns)):
                try:
                    for key, value in DICT.items():
                        if key in table.cell(row ,col).text:
                            table.cell(row ,col).text = table.cell(row ,col).text.replace(key, value)
                except Exception as e:
                    continue
                

    # paragraphs
    for para in document.paragraphs:       
        for i in range(len(para.runs)):
            for key, value in DICT.items():                
                if key in para.runs[i].text:
                    para.runs[i].text = para.runs[i].text.replace(key, value)

    return document 
    






if __name__ == '__main__':
	main()